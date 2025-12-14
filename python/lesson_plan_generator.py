#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
指導案ジェネレーター（シンプル版）
単一YAMLコンテンツからHTML形式またはWord形式の指導案を生成する
"""

import yaml
import io
import base64
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class LessonPlanGenerator:
    def __init__(self, yaml_content: str):
        """YAMLコンテンツから初期化"""
        self.data = yaml.safe_load(yaml_content)

    def generate_html(self) -> str:
        """HTML文字列を生成して返す"""
        return self._build_html()

    def generate_docx_bytes(self) -> bytes:
        """Word文書をバイト列として生成"""
        doc = self._build_docx()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_docx_base64(self) -> str:
        """Word文書をBase64エンコードして返す"""
        docx_bytes = self.generate_docx_bytes()
        return base64.b64encode(docx_bytes).decode('utf-8')

    def _set_cell_shading(self, cell, color: str):
        """セルの背景色を設定"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _build_docx(self):
        """Word文書を構築"""
        d = self.data
        doc = Document()
        
        # ページ設定
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # タイトル
        heading = doc.add_heading(f"{d.get('教科', '')}科 学習指導案", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ヘッダー表
        header_table = doc.add_table(rows=5, cols=2)
        header_table.style = 'Table Grid'
        
        header_data = [
            ('日　時', d.get('日時', '')),
            ('学校名', d.get('学校名', '')),
            ('対　象', d.get('対象', '')),
            ('会　場', d.get('会場', '')),
            ('授業者', d.get('授業者', '')),
        ]
        for i, (label, value) in enumerate(header_data):
            header_table.cell(i, 0).text = label
            header_table.cell(i, 1).text = str(value)
            self._set_cell_shading(header_table.cell(i, 0), 'F5F5F5')
        
        doc.add_paragraph()
        
        # 1. 単元名
        doc.add_heading('１　単元名', level=1)
        p = doc.add_paragraph()
        run = p.add_run(d.get('単元名', ''))
        run.bold = True
        p.add_run(f"（{d.get('使用教科書', '')}）")
        
        # 2. 本時の目標
        doc.add_heading('２　本時の目標', level=1)
        goals = d.get('本時の目標', d.get('目標', []))
        for goal in goals:
            if goal:
                doc.add_paragraph(goal, style='List Bullet')
        
        # 3. 本時の展開
        doc.add_heading('３　本時の展開', level=1)
        flow = d.get('展開', d.get('授業展開', {}))
        
        if flow:
            flow_table = doc.add_table(rows=1 + len(flow), cols=3)
            flow_table.style = 'Table Grid'
            
            # ヘッダー行
            headers = ['時間', '○学習内容　・学習活動', '指導上の留意点']
            for i, h in enumerate(headers):
                cell = flow_table.cell(0, i)
                cell.text = h
                self._set_cell_shading(cell, 'E8E8E8')
            
            # データ行
            row_idx = 1
            for phase_name, phase in flow.items():
                if not phase:
                    continue
                
                time_str = f"{phase_name}\n({phase.get('時間', '')}分)"
                
                content = []
                for c in phase.get('学習内容', []):
                    if c:
                        content.append(f"○{c}")
                for a in phase.get('学習活動', []):
                    if a:
                        content.append(f"・{a}")
                
                notes = '\n'.join([f"・{n}" for n in phase.get('留意点', []) if n])
                
                flow_table.cell(row_idx, 0).text = time_str
                flow_table.cell(row_idx, 1).text = '\n'.join(content)
                flow_table.cell(row_idx, 2).text = notes
                row_idx += 1
        
        # 4. 本時の評価
        evaluations = d.get('評価', d.get('本時の評価', []))
        if evaluations:
            doc.add_heading('４　本時の評価', level=1)
            for e in evaluations:
                if e:
                    text = e.get('規準', e) if isinstance(e, dict) else e
                    doc.add_paragraph(text, style='List Bullet')
        
        return doc

    def _build_html(self):
        d = self.data
        return f"""<!DOCTYPE html>
<html lang="ja">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{d.get('教科', '')}科 学習指導案</title>
    <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
    <style>
        @media print {{
            @page {{ size: A4; margin: 20mm; }}
            .page-break {{ page-break-after: always; }}
        }}
        body {{ 
            font-family: 'Hiragino Mincho ProN', 'Yu Mincho', serif; 
            line-height: 1.6; 
            max-width: 210mm; 
            margin: 0 auto; 
            padding: 20px;
            color: #333;
        }}
        
        h1 {{ text-align: center; font-size: 22pt; margin-bottom: 30px; border-bottom: 3px double #000; padding-bottom: 10px; }}
        h2 {{ font-size: 14pt; margin-top: 25px; border-left: 4px solid #3b82f6; padding-left: 10px; background: #f0f8ff; padding: 8px 10px; }}
        h3 {{ font-size: 12pt; margin-top: 15px; }}
        
        .header-table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
        .header-table td, .header-table th {{ border: 1px solid #333; padding: 8px 12px; }}
        .header-table th {{ background: #f5f5f5; width: 100px; text-align: left; }}
        
        .section {{ margin: 20px 0; }}
        .section ul {{ margin: 10px 0; padding-left: 25px; }}
        .section li {{ margin: 5px 0; }}
        
        .flow-table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        .flow-table th, .flow-table td {{ border: 1px solid #333; padding: 10px; vertical-align: top; }}
        .flow-table th {{ background: #e8e8e8; text-align: center; }}
        .flow-table td:first-child {{ width: 80px; text-align: center; font-weight: bold; }}
        
        .activity {{ margin: 5px 0; }}
        .activity-content {{ color: #000; }}
        .activity-action {{ color: #555; margin-left: 15px; }}
        
        .goals {{ background: #fffde7; padding: 15px; border-radius: 5px; border-left: 4px solid #ffc107; }}
        .evaluation {{ background: #e8f5e9; padding: 15px; border-radius: 5px; border-left: 4px solid #4caf50; }}
    </style>
</head>
<body>
    <h1>{d.get('教科', '')}科 学習指導案</h1>
    
    {self._create_header()}
    {self._create_unit_info()}
    {self._create_goals()}
    {self._create_flow()}
    {self._create_evaluation()}
</body>
</html>"""

    def _create_header(self):
        d = self.data
        return f"""
    <table class="header-table">
        <tr><th>日　時</th><td>{d.get('日時', '')}</td></tr>
        <tr><th>学校名</th><td>{d.get('学校名', '')}</td></tr>
        <tr><th>対　象</th><td>{d.get('対象', '')}</td></tr>
        <tr><th>会　場</th><td>{d.get('会場', '')}</td></tr>
        <tr><th>授業者</th><td>{d.get('授業者', '')}</td></tr>
    </table>"""

    def _create_unit_info(self):
        d = self.data
        return f"""
    <h2>１　単元名</h2>
    <div class="section">
        <strong>{d.get('単元名', '')}</strong>
        （{d.get('使用教科書', '')}）
    </div>"""

    def _create_goals(self):
        d = self.data
        goals = d.get('本時の目標', [])
        if not goals:
            goals = d.get('目標', [])
        
        goals_html = "\n".join([f"<li>{g}</li>" for g in goals if g])
        
        return f"""
    <h2>２　本時の目標</h2>
    <div class="section goals">
        <ul>{goals_html}</ul>
    </div>"""

    def _create_flow(self):
        d = self.data
        flow = d.get('展開', d.get('授業展開', {}))
        
        if not flow:
            return ""
        
        rows_html = ""
        for phase_name, phase in flow.items():
            if not phase:
                continue
            
            time = phase.get('時間', '')
            
            # 学習内容・活動
            activities = []
            for c in phase.get('学習内容', []):
                if c:
                    activities.append(f'<div class="activity"><span class="activity-content">○ {c}</span></div>')
            for a in phase.get('学習活動', []):
                if a:
                    activities.append(f'<div class="activity"><span class="activity-action">・ {a}</span></div>')
            activities_html = "\n".join(activities)
            
            # 留意点
            notes = phase.get('留意点', [])
            notes_html = "\n".join([f"・{n}" for n in notes if n])
            
            rows_html += f"""
        <tr>
            <td>{phase_name}<br>({time}分)</td>
            <td>{activities_html}</td>
            <td>{notes_html}</td>
        </tr>"""
        
        return f"""
    <h2>３　本時の展開</h2>
    <table class="flow-table">
        <tr>
            <th>時間</th>
            <th>○学習内容　・学習活動</th>
            <th>指導上の留意点</th>
        </tr>
        {rows_html}
    </table>"""

    def _create_evaluation(self):
        d = self.data
        evaluations = d.get('評価', d.get('本時の評価', []))
        
        if not evaluations:
            return ""
        
        if isinstance(evaluations, list):
            evals_html = "\n".join([f"<li>{e.get('規準', e) if isinstance(e, dict) else e}</li>" for e in evaluations if e])
        else:
            evals_html = f"<li>{evaluations}</li>"
        
        return f"""
    <h2>４　本時の評価</h2>
    <div class="section evaluation">
        <ul>{evals_html}</ul>
    </div>"""


def generate_lesson_plan_html(yaml_content: str) -> str:
    """YAML文字列からHTML文字列を生成"""
    generator = LessonPlanGenerator(yaml_content)
    return generator.generate_html()


def generate_lesson_plan_docx_base64(yaml_content: str) -> str:
    """YAML文字列からWord文書をBase64で生成"""
    generator = LessonPlanGenerator(yaml_content)
    return generator.generate_docx_base64()

